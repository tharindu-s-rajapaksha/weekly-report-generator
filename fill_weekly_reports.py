import json
import os
import re
from typing import Dict, List, Optional

from docx import Document, shared


def load_data(path: Optional[str]) -> Dict:
    if path and os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)

def set_cell_text(cell, text: str) -> None:
    """Set cell text by clearing and adding new content"""
    # Clear existing paragraphs
    for paragraph in cell.paragraphs:
        paragraph.clear()
    # Add new text
    if cell.paragraphs:
        cell.paragraphs[0].add_run(str(text) if text is not None else "")
    else:
        cell.add_paragraph().add_run(str(text) if text is not None else "")


def fill_week_header_info(doc: Document, data: Dict) -> None:
    """Fill the week number, week ending, and training mode in the header area"""
    if not doc.tables:
        return
    
    table = doc.tables[0]  # Main table
    
    # The first row contains week ending info and training mode
    if len(table.rows) > 0:
        first_row = table.rows[0]
        
        # Update week ending (first few cells should contain this)
        week_ending = data.get("week_ending", "")
        if len(first_row.cells) > 0 and "FOR THE WEEK ENDING" in first_row.cells[0].text:
            set_cell_label_value(first_row.cells[0], "FOR THE WEEK ENDING", week_ending)
        
        # Update training mode (should be in the rightmost cell)
        training_mode = data.get("training_mode", "")
        if len(first_row.cells) > 3 and "TRAINING MODE" in first_row.cells[3].text:
            set_cell_label_value(first_row.cells[3], "TRAINING MODE", training_mode)
            
def set_week_no(doc: Document, week_no: str) -> None:
    """Find the WEEK NO label in the document and set the full text with the number, preserving previous styling (bold)."""
    pat = re.compile(r"WEEK\s*NO", re.I)
    for p in doc.paragraphs:
        if pat.search(p.text):
            # Find the run containing "WEEK NO" and modify its text, preserving styling
            for run in p.runs:
                if pat.search(run.text):
                    # Preserve the run's styling by modifying its text
                    run.text = f"WEEK NO: {week_no}"
                    return

def set_cell_label_value(cell, label: str, value: str, label_bold: bool = True) -> None:
    """Clear a cell and set two runs: a bold label and a normal value (on new line)."""
    # Clear existing paragraphs
    for paragraph in cell.paragraphs:
        paragraph.clear()
    # Use a single paragraph and add runs
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    # Add label run
    if label:
        r_label = p.add_run(str(label))
        r_label.bold = bool(label_bold)
    # Add newline then value
    if value is not None:
        p.add_run("\n")
        r_val = p.add_run(str(value))
        r_val.bold = False


def fill_daily_activities(doc: Document, activities: List[Dict]) -> None:
    """Fill the daily activities in the main table"""
    if not doc.tables or not activities:
        return
    
    table = doc.tables[0]
    
    # Create mapping from day to activity data
    activity_map = {}
    for activity in activities:
        day_key = activity["day"].upper().strip()
        activity_map[day_key] = activity
    
    # Start from row 2 (skip header rows)
    for row_idx in range(2, len(table.rows)):
        row = table.rows[row_idx]
        if len(row.cells) < 3:
            continue
            
        # Get the day from the first cell
        day_cell_text = row.cells[0].text.strip().upper()
        
        # Extract day name from vertical text (might contain newlines)
        day_name = ""
        for day in ["MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY", "SATURDAY", "SUNDAY"]:
            if day in day_cell_text:
                day_name = day
                break
        
        if day_name and day_name in activity_map:
            activity = activity_map[day_name]
            
            # Fill date (column 1)
            if len(row.cells) > 1:
                set_cell_text(row.cells[1], activity.get("date", ""))
            
            # Fill description (column 2 - the large column)
            if len(row.cells) > 2:
                set_cell_text(row.cells[2], activity.get("description", ""))


def fill_details_section(doc: Document, data: Dict) -> None:
    """Fill the details and notes section at the bottom"""
    if len(doc.tables) < 2:
        return
    
    table = doc.tables[1]  # Second table for details
    
    # Fill the main details text area (first large cell)
    details_notes = data.get("details_notes", "")
    if len(table.rows) > 1 and len(table.rows[1].cells) > 0:
        set_cell_text(table.rows[1].cells[0], details_notes)
        # Fill trainee signature
        # Look for cell containing "SIGNATURE OF TRAINEE"
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if "SIGNATURE OF TRAINEE" in cell.text:
                    # Bold the label cell itself, keep label and leave value empty there
                    set_cell_label_value(cell, "SIGNATURE OF TRAINEE", "")
                    # Put the actual signature image in the adjacent cell if available
                    if cell_idx + 1 < len(row.cells):
                        # Clear the adjacent cell
                        for paragraph in row.cells[cell_idx + 1].paragraphs:
                            paragraph.clear()
                        # Add image to the adjacent cell
                        if row.cells[cell_idx + 1].paragraphs:
                            p = row.cells[cell_idx + 1].paragraphs[0]
                        else:
                            p = row.cells[cell_idx + 1].add_paragraph()
                        run = p.add_run()
                        run.add_picture("data/signature.png", width=shared.Inches(1.5), height=shared.Inches(0.5))
                    else:
                        # If no adjacent cell, add image below label in same cell
                        # Clear existing paragraphs
                        for paragraph in cell.paragraphs:
                            paragraph.clear()
                        # Add label run
                        p = cell.add_paragraph()
                        r_label = p.add_run("SIGNATURE OF TRAINEE")
                        r_label.bold = True
                        # Add newline then image
                        p.add_run("\n")
                        run = p.add_run()
                        run.add_picture("data/signature.png", width=shared.Inches(1.5), height=shared.Inches(0.5))
                    break

    # Fill engineer remarks
    engineer_remarks = data.get("engineer_remarks", "")
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if "REMARKS AND CERTIFICATION" in cell.text:
                # Use the row below this header
                if row_idx + 1 < len(table.rows):
                    set_cell_text(table.rows[row_idx + 1].cells[0], engineer_remarks)
                break
    
    # Fill date and engineer signature
    engineer_date = data.get("engineer_date", "")
    engineer_designation = data.get("engineer_designation_signature", "")
    
    # Look for the last row which should have DATE and DESIGNATION fields
    if len(table.rows) > 0:
        last_row = table.rows[-1]
        for cell_idx, cell in enumerate(last_row.cells):
            if "DATE:" in cell.text:
                # Make the DATE label bold and put the date value below
                set_cell_label_value(cell, "DATE:", engineer_date)
            elif "DESIGNATION AND SIGNATURE" in cell.text:
                # set label bold and value
                set_cell_label_value(cell, "DESIGNATION AND SIGNATURE", engineer_designation)


def main():
    template_path = "data/Daily Report Template.docx"
    data_path = "data/weekly_data.json"

    if not os.path.exists(template_path):
        raise SystemExit(f"Template not found: {template_path}")

    all_data = load_data(data_path)

    for week_data in all_data:
        output_path = f"Weekly Reports/Daily Report {week_data.get('week_no', '')}.docx"

        if not os.path.exists("Weekly Reports"):
            os.makedirs("Weekly Reports")

        doc = Document(template_path)
        # Fill different sections
        fill_week_header_info(doc, week_data)
        # Set WEEK NO dynamically (bold label)
        week_no = week_data.get("week_no", "")
        if week_no:
            set_week_no(doc, week_no)
        
        weekly_activities = week_data.get("weekly_activities", [])
        if weekly_activities:
            fill_daily_activities(doc, weekly_activities)
        
        fill_details_section(doc, week_data)

        # Save output
        doc.save(output_path)
        print(f"Saved: {output_path}")


if __name__ == "__main__":
    main()
